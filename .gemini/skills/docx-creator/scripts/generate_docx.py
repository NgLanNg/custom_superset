#!/usr/bin/env python3
"""
Generate a styled Word document from Markdown content and JSON styles.
"""
import sys
import re
import json
import argparse
from datetime import datetime

import os

try:
 from docx import Document
 from docx.shared import Inches, Pt, RGBColor
 from docx.enum.text import WD_ALIGN_PARAGRAPH
 from docx.enum.style import WD_STYLE_TYPE
 from docx.oxml.ns import nsdecls
 from docx.oxml import parse_xml
except ImportError:
 print("Error: python-docx is not installed. Please run: pip install python-docx")
 sys.exit(1)

def hex_to_rgb(hex_str):
 hex_str = hex_str.lstrip('#')
 return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))

def set_cell_bg(cell, color_hex):
 """Set background color for a table cell using XML manipulation."""
 if not color_hex: return
 shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
 cell._tc.get_or_add_tcPr().append(shading_elm)

def apply_style_config(doc, config):
 """Apply styles from JSON config to the document."""
 styles = doc.styles

 # Document Margins
 if "document" in config:
 section = doc.sections[0]
 margins = config["document"]
 if "margin_top_inches" in margins: section.top_margin = Inches(margins["margin_top_inches"])
 if "margin_bottom_inches" in margins: section.bottom_margin = Inches(margins["margin_bottom_inches"])
 if "margin_left_inches" in margins: section.left_margin = Inches(margins["margin_left_inches"])
 if "margin_right_inches" in margins: section.right_margin = Inches(margins["margin_right_inches"])

 # Font Defaults
 default_font = config.get("fonts", {}).get("default", "Calibri")
 if 'Normal' in styles:
 doc.styles['Normal'].font.name = default_font

 # Custom Styles
 style_configs = config.get("styles", {})

 for style_name, rules in style_configs.items():
 # Get or create style
 if style_name in styles:
 style = styles[style_name]
 else:
 try:
 style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
 except:
 continue

 # Font properties
 if "font" in rules: style.font.name = rules["font"]
 if "size_pt" in rules: style.font.size = Pt(rules["size_pt"])
 if "color_hex" in rules: style.font.color.rgb = hex_to_rgb(rules["color_hex"])
 if "bold" in rules: style.font.bold = rules["bold"]
 if "italic" in rules: style.font.italic = rules["italic"]
 if "uppercase" in rules: style.font.all_caps = rules["uppercase"]

 # Paragraph formatting
 if "alignment" in rules:
 align_map = {"LEFT": WD_ALIGN_PARAGRAPH.LEFT, "CENTER": WD_ALIGN_PARAGRAPH.CENTER, "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT, "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY}
 if rules["alignment"] in align_map:
 style.paragraph_format.alignment = align_map[rules["alignment"]]

 if "space_before_pt" in rules: style.paragraph_format.space_before = Pt(rules["space_before_pt"])
 if "space_after_pt" in rules: style.paragraph_format.space_after = Pt(rules["space_after_pt"])


def parse_markdown_to_docx(md_path, json_style_path, output_path):
 doc = Document()

 # Load Styles
 try:
 with open(json_style_path, 'r') as f:
 style_config = json.load(f)
 except FileNotFoundError:
 print(f"Warning: Style file '{json_style_path}' not found. Using defaults.")
 style_config = {}

 apply_style_config(doc, style_config)

 # --- Read Markdown ---
 try:
 with open(md_path, 'r', encoding='utf-8') as f:
 lines = f.readlines()
 except FileNotFoundError:
 print(f"Error: Markdown file '{md_path}' not found.")
 sys.exit(1)

 # --- Processing Loop ---
 table_mode = False
 table_data = []

 for line in lines:
 line = line.strip()

 if not line:
 if table_mode:
 create_table(doc, table_data, style_config)
 table_mode = False
 table_data = []
 continue

 # Title Page (H1 at start)
 if line.startswith('# '):
 title = line[2:]
 create_title_page(doc, title, style_config)

 # Headings
 elif line.startswith('## '):
 h = doc.add_heading(line[3:], level=1)
 elif line.startswith('### '):
 doc.add_heading(line[4:], level=2)

 # Lists
 elif line.startswith('- '):
 p = doc.add_paragraph(style='List Bullet')
 process_inline_formatting(p, line[2:])

 # Tables
 elif line.startswith('|'):
 table_mode = True
 cells = [c.strip() for c in line.strip('|').split('|')]
 if '---' in cells[0]: continue
 table_data.append(cells)

 # Key-Value Metadata
 elif line.startswith('**') and ':**' in line:
 p = doc.add_paragraph()
 p.alignment = WD_ALIGN_PARAGRAPH.CENTER
 parts = line.split(':**', 1)
 key = parts[0].replace('**', '') + ':'
 value = parts[1].strip()

 run = p.add_run(key)
 run.bold = True
 run.font.size = Pt(12)
 p.add_run(f" {value}")

 # Images: ![caption](path)
 elif line.startswith('![') and '](' in line:
 if table_mode:
 create_table(doc, table_data, style_config)
 table_mode = False
 table_data = []

 # Parse ![caption](path)
 match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
 if match:
 caption = match.group(1)
 img_path = match.group(2)
 add_image_with_caption(doc, img_path, caption, style_config)

 # Normal Text
 else:
 if table_mode:
 create_table(doc, table_data, style_config)
 table_mode = False
 table_data = []

 p = doc.add_paragraph(style='Normal')
 process_inline_formatting(p, line)

 if table_mode and table_data:
 create_table(doc, table_data, style_config)

 # --- Save ---
 try:
 doc.save(output_path)
 print(f"Successfully created document at: {output_path}")
 except PermissionError:
 print(f"Error: Could not save to '{output_path}'. File might be open.")

def create_title_page(doc, title, config):
 for _ in range(4): doc.add_paragraph()

 style = 'Title' if 'Title' in doc.styles else 'Heading 1'
 p = doc.add_paragraph(title, style=style)

 doc.add_paragraph() # Spacing

def add_image_with_caption(doc, img_path, caption, config):
 """Add an image with a centered caption below it."""
 # Resolve relative paths
 if not os.path.isabs(img_path):
 # Try relative to current working directory
 if not os.path.exists(img_path):
 print(f"Warning: Image not found at '{img_path}'")
 # Add placeholder text
 p = doc.add_paragraph(f"[Image not found: {img_path}]")
 p.alignment = WD_ALIGN_PARAGRAPH.CENTER
 return

 try:
 # Add the image centered
 p = doc.add_paragraph()
 p.alignment = WD_ALIGN_PARAGRAPH.CENTER
 run = p.add_run()

 # Get image width from config or use default
 img_width = config.get("images", {}).get("width_inches", 5.5)
 run.add_picture(img_path, width=Inches(img_width))

 # Add caption below
 if caption:
 cap_p = doc.add_paragraph()
 cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
 cap_run = cap_p.add_run(caption)
 cap_run.italic = True
 cap_run.font.size = Pt(10)
 cap_run.font.color.rgb = RGBColor(100, 100, 100)

 # Add spacing after
 doc.add_paragraph()

 except Exception as e:
 print(f"Error adding image '{img_path}': {e}")
 p = doc.add_paragraph(f"[Error loading image: {img_path}]")
 p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_table(doc, data, config):
 if not data: return

 rows = len(data)
 cols = len(data[0])
 table = doc.add_table(rows=rows, cols=cols)
 table.style = 'Table Grid'

 # Header Config
 header_bg = config.get("styles", {}).get("TableHeader", {}).get("bg_color_hex", "2E74B5")
 header_text_color = config.get("styles", {}).get("TableHeader", {}).get("color_hex", "FFFFFF")

 for r_idx, row_data in enumerate(data):
 row = table.rows[r_idx]
 for c_idx, cell_text in enumerate(row_data):
 if c_idx < len(row.cells):
 cell = row.cells[c_idx]

 # Clear default text and use inline formatting
 cell.text = ""
 if cell.paragraphs:
 p = cell.paragraphs[0]
 else:
 p = cell.add_paragraph()

 # Apply Header Styling
 if r_idx == 0:
 set_cell_bg(cell, header_bg)
 # For headers, just bold the text
 run = p.add_run(cell_text)
 run.font.bold = True
 if header_text_color:
 run.font.color.rgb = hex_to_rgb(header_text_color)
 else:
 # For data rows, process inline formatting
 process_inline_formatting(p, cell_text)

def process_inline_formatting(paragraph, text):
 """Process inline markdown formatting (bold, italic) in text."""
 # Handle both **bold** and *italic*
 parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
 for part in parts:
 if part.startswith('**') and part.endswith('**') and len(part) > 4:
 paragraph.add_run(part[2:-2]).bold = True
 elif part.startswith('*') and part.endswith('*') and len(part) > 2 and not part.startswith('**'):
 paragraph.add_run(part[1:-1]).italic = True
 else:
 paragraph.add_run(part)

def main():
 parser = argparse.ArgumentParser(description="Generate a styled Word document from Markdown.")
 parser.add_argument("--input", "-i", required=True, help="Input Markdown file")
 parser.add_argument("--style", "-s", required=True, help="Style JSON file")
 parser.add_argument("--output", "-o", default="output.docx", help="Output DOCX file")

 args = parser.parse_args()

 parse_markdown_to_docx(args.input, args.style, args.output)

if __name__ == "__main__":
 main()
