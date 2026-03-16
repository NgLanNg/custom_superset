#!/usr/bin/env python3
"""
Generate business documents (BRD, Proposal, etc.) from templates.

Usage:
 python generate_doc.py --type brd --output output.md
 python generate_doc.py --type proposal --format docx --output output.docx
"""

import argparse
import os
import sys
from pathlib import Path
from datetime import datetime

# Try to import docx-creator functionality if available
try:
 from docx import Document
 DOCX_AVAILABLE = True
except ImportError:
 DOCX_AVAILABLE = False


def get_template_path(doc_type: str) -> Path:
 """Get the template file path for a given document type."""
 script_dir = Path(__file__).parent
 skill_dir = script_dir.parent
 template_dir = skill_dir / "assets" / "templates"

 template_map = {
 "brd": "brd_template.md",
 "proposal": "proposal_template.md",
 "frs": "frs_template.md",
 "technical": "technical_spec_template.md"
 }

 template_file = template_map.get(doc_type.lower())
 if not template_file:
 raise ValueError(f"Unknown document type: {doc_type}")

 return template_dir / template_file


def load_template(template_path: Path) -> str:
 """Load template content from file."""
 if not template_path.exists():
 raise FileNotFoundError(f"Template not found: {template_path}")

 with open(template_path, 'r', encoding='utf-8') as f:
 return f.read()


def fill_template(template_content: str, replacements: dict) -> str:
 """Replace placeholders in template with actual values."""
 content = template_content

 # Add default replacements
 defaults = {
 "[YYYY-MM-DD]": datetime.now().strftime("%Y-%m-%d"),
 "[Date]": datetime.now().strftime("%Y-%m-%d"),
 }

 all_replacements = {**defaults, **replacements}

 for placeholder, value in all_replacements.items():
 content = content.replace(placeholder, value)

 return content


def save_markdown(content: str, output_path: Path):
 """Save content as Markdown file."""
 with open(output_path, 'w', encoding='utf-8') as f:
 f.write(content)
 print(f"✓ Markdown document created: {output_path}")


def save_docx(content: str, output_path: Path):
 """Save content as DOCX file (requires python-docx)."""
 if not DOCX_AVAILABLE:
 print("ERROR: python-docx not installed. Install with: pip install python-docx", file=sys.stderr)
 sys.exit(1)

 doc = Document()

 # Simple conversion: split by lines and add paragraphs
 # For production use, integrate with docx-creator skill for better formatting
 lines = content.split('\n')

 for line in lines:
 # Handle headers
 if line.startswith('# '):
 doc.add_heading(line[2:], level=1)
 elif line.startswith('## '):
 doc.add_heading(line[3:], level=2)
 elif line.startswith('### '):
 doc.add_heading(line[4:], level=3)
 else:
 doc.add_paragraph(line)

 doc.save(output_path)
 print(f"✓ DOCX document created: {output_path}")


def main():
 parser = argparse.ArgumentParser(
 description="Generate business documents from templates"
 )
 parser.add_argument(
 "--type",
 required=True,
 choices=["brd", "proposal", "frs", "technical"],
 help="Document type to generate"
 )
 parser.add_argument(
 "--format",
 default="md",
 choices=["md", "docx"],
 help="Output format (default: md)"
 )
 parser.add_argument(
 "--output",
 required=True,
 help="Output file path"
 )
 parser.add_argument(
 "--project-name",
 help="Project name to fill in template"
 )
 parser.add_argument(
 "--prepared-by",
 help="Author name"
 )

 args = parser.parse_args()

 try:
 # Load template
 template_path = get_template_path(args.type)
 template_content = load_template(template_path)

 # Prepare replacements
 replacements = {}
 if args.project_name:
 replacements["[Project Name]"] = args.project_name
 if args.prepared_by:
 replacements["[Your Name]"] = args.prepared_by

 # Fill template
 content = fill_template(template_content, replacements)

 # Save output
 output_path = Path(args.output)

 if args.format == "md":
 save_markdown(content, output_path)
 elif args.format == "docx":
 save_docx(content, output_path)

 except Exception as e:
 print(f"ERROR: {e}", file=sys.stderr)
 sys.exit(1)


if __name__ == "__main__":
 main()
